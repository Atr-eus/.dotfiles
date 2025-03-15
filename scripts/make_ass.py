#!/usr/bin/python3

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys


def insert_content(doc, content, font_sz, bold, gap):
    if content.endswith((".py", ".cpp", ".c")):
        try:
            with open(content, "r") as file:
                content = file.read()
        except FileNotFoundError:
            print(f"File not found: {content}")
            return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(content)
    run.font.name = "Courier New"
    run.font.size = Pt(font_sz)
    run.bold = bold
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("\n" * gap)


def main():
    if len(sys.argv) < 2:
        print("Usage: ./make_ass.py example.docx")
        return

    output_path = sys.argv[1]
    docx = Document()

    while True:
        content = input("Enter content ['q' to quit]: ")
        if content == "q":
            break

        font_sz = int(input("Font size? "))
        bold = input("Bold? [y/n] ").strip().lower() == "y"
        gap = int(input("Gap? "))

        insert_content(docx, content, font_sz, bold, gap)

    docx.save(output_path)
    print(f"Document saved as {output_path}")


if __name__ == "__main__":
    main()
